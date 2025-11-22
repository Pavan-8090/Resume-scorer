import io
import os
import json
from typing import Dict, List
import fitz  # PyMuPDF
import docx2txt
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import re
from huggingface_hub import HfApi, snapshot_download
from dotenv import load_dotenv
import requests

load_dotenv()

# Try to import OpenAI
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("WARNING: OpenAI not installed. Install with: pip install openai")

# Try to import PyResparser and spaCy
try:
    from pyresparser import ResumeParser
    PYRESPARSER_AVAILABLE = True
except ImportError:
    PYRESPARSER_AVAILABLE = False
    print("WARNING: PyResparser not available. Install with: pip install pyresparser")

try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except OSError:
        print("WARNING: spaCy model not found. Run: python -m spacy download en_core_web_sm")
        SPACY_AVAILABLE = False
except ImportError:
    SPACY_AVAILABLE = False
    print("WARNING: spaCy not available. Install with: pip install spacy")

# Load SentenceTransformer model
_embedder = None
_hf_api = None

def get_hf_api():
    """Get Hugging Face API client"""
    global _hf_api
    if _hf_api is None:
        hf_token = os.getenv("HF_TOKEN")
        if hf_token:
            _hf_api = HfApi(token=hf_token)
            print("SUCCESS: Initialized Hugging Face Hub API")
        else:
            print("WARNING: HF_TOKEN not set. Hugging Face Hub features disabled.")
    return _hf_api

def load_model_from_hf(repo_id: str = None):
    """Load model from Hugging Face Hub or use default"""
    if repo_id:
        try:
            hf_api = get_hf_api()
            if hf_api:
                print(f"Loading model from Hugging Face Hub: {repo_id}")
                # Download model if not cached
                model_path = snapshot_download(
                    repo_id=repo_id,
                    token=os.getenv("HF_TOKEN"),
                    cache_dir=None  # Use default cache
                )
                return SentenceTransformer(model_path)
            else:
                print(f"HF_TOKEN not available, trying to load {repo_id} without auth...")
                return SentenceTransformer(repo_id)
        except Exception as e:
            print(f"Failed to load model from HF Hub {repo_id}: {e}")
            print("Falling back to default model...")
    
    # Default model
    return SentenceTransformer('all-MiniLM-L6-v2')

def get_embedder():
    """Get or initialize the embedding model"""
    global _embedder
    if _embedder is None:
        try:
            # Try to load custom model from HF Hub if specified
            custom_model = os.getenv("HF_MODEL_ID", "all-MiniLM-L6-v2")
            
            if custom_model.startswith("PavanRathodR/") or "/" in custom_model:
                # Custom model from Hugging Face Hub
                _embedder = load_model_from_hf(custom_model)
            else:
                # Default model
                _embedder = SentenceTransformer(custom_model)
            
            print(f"SUCCESS: Loaded SentenceTransformer model: {custom_model}")
        except Exception as e:
            print(f"WARNING: Could not load SentenceTransformer: {e}")
            _embedder = None
    return _embedder

def upload_model_to_hf(folder_path: str, repo_id: str):
    """Upload a trained model to Hugging Face Hub"""
    try:
        hf_api = get_hf_api()
        if not hf_api:
            raise ValueError("HF_TOKEN not set. Cannot upload to Hugging Face Hub.")
        
        if not os.path.exists(folder_path):
            raise ValueError(f"Model folder not found: {folder_path}")
        
        print(f"Uploading model from {folder_path} to {repo_id}...")
        hf_api.upload_folder(
            folder_path=folder_path,
            repo_id=repo_id,
            repo_type="model",
        )
        print(f"SUCCESS: Model uploaded to {repo_id}")
        return True
    except Exception as e:
        print(f"ERROR: Failed to upload model: {e}")
        return False

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX using docx2txt"""
    try:
        text = docx2txt.process(io.BytesIO(file_content))
        return text
    except Exception as e:
        raise Exception(f"Failed to parse DOCX: {str(e)}")

def extract_resume_text(file_content: bytes, file_name: str) -> str:
    """Extract text from resume file"""
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_ext in ['docx', 'doc']:
        return extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def extract_skills_with_spacy(resume_text: str) -> List[str]:
    """Extract skills using spaCy"""
    if not SPACY_AVAILABLE:
        return []
    
    doc = nlp(resume_text)
    skills = []
    
    # Common skill keywords
    skill_keywords = [
        'javascript', 'python', 'java', 'react', 'node.js', 'sql', 'html', 'css',
        'typescript', 'angular', 'vue', 'express', 'django', 'flask', 'spring',
        'leadership', 'communication', 'project management', 'teamwork', 'problem solving',
        'marketing', 'sales', 'analytics', 'design', 'writing', 'analysis', 'agile',
        'devops', 'aws', 'docker', 'kubernetes', 'git', 'ci/cd', 'machine learning',
        'data science', 'ai', 'tensorflow', 'pytorch', 'postgresql', 'mongodb'
    ]
    
    resume_lower = resume_text.lower()
    for keyword in skill_keywords:
        if keyword in resume_lower:
            skills.append(keyword.title())
    
    return list(set(skills))[:10]

def parse_resume_with_pyresparser(file_content: bytes, file_name: str) -> Dict:
    """Parse resume using PyResparser"""
    if not PYRESPARSER_AVAILABLE:
        return {}
    
    try:
        # Save file temporarily
        import tempfile
        import os
        
        file_ext = file_name.split('.')[-1] if '.' in file_name else 'pdf'
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_ext}')
        temp_file.write(file_content)
        temp_file.close()
        
        # Parse resume
        data = ResumeParser(temp_file.name).get_extracted_data()
        
        # Clean up
        os.unlink(temp_file.name)
        
        return data
    except Exception as e:
        print(f"PyResparser error: {e}")
        return {}

def calculate_semantic_similarity(text1: str, text2: str) -> float:
    """Calculate semantic similarity using SentenceTransformer"""
    embedder = get_embedder()
    if embedder is None:
        return calculate_keyword_match(text1, text2)
    
    try:
        # Truncate to avoid token limits
        text1_truncated = text1[:2000] if len(text1) > 2000 else text1
        text2_truncated = text2[:2000] if len(text2) > 2000 else text2
        
        embeddings = embedder.encode([text1_truncated, text2_truncated])
        similarity = cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]
        return float(similarity)
    except Exception as e:
        print(f"Error in semantic similarity: {e}")
        return calculate_keyword_match(text1, text2)

def calculate_keyword_match(resume_text: str, job_description: str) -> float:
    """Fallback keyword matching - improved accuracy"""
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    
    common_skills = [
        'javascript', 'python', 'java', 'react', 'node.js', 'sql', 'html', 'css',
        'typescript', 'angular', 'vue', 'express', 'django', 'flask', 'spring',
        'leadership', 'communication', 'project management', 'teamwork', 'problem solving',
        'marketing', 'sales', 'analytics', 'design', 'writing', 'analysis', 'agile',
        'devops', 'aws', 'docker', 'kubernetes', 'git', 'ci/cd', 'machine learning',
        'data science', 'ai', 'tensorflow', 'pytorch', 'postgresql', 'mongodb', 'redis',
        'elasticsearch', 'kafka', 'rabbitmq', 'nginx', 'apache', 'linux', 'terraform',
        'ansible', 'jenkins', 'github actions', 'graphql', 'rest api', 'microservices'
    ]
    
    # Count matched skills
    matched_skills = [skill for skill in common_skills if skill in job_lower and skill in resume_lower]
    required_skills = [skill for skill in common_skills if skill in job_lower]
    
    if len(required_skills) == 0:
        # No specific skills mentioned, use general text similarity
        resume_words = set(resume_lower.split())
        job_words = set(job_lower.split())
        common_words = resume_words.intersection(job_words)
        if len(job_words) > 0:
            return min(0.7, len(common_words) / len(job_words))
        return 0.5
    
    # Calculate skill match percentage
    skill_match_ratio = len(matched_skills) / len(required_skills) if required_skills else 0
    
    # Base score from skill matching (0-0.8 range)
    base_score = min(0.8, skill_match_ratio * 0.8)
    
    # Add bonus for having extra relevant skills (0-0.2 range)
    extra_skills = [skill for skill in common_skills if skill in resume_lower and skill not in required_skills]
    bonus = min(0.2, len(extra_skills) * 0.02)
    
    final_score = base_score + bonus
    return min(1.0, max(0.0, final_score))

def extract_job_required_skills(job_description: str) -> List[str]:
    """Extract skills mentioned in job description - comprehensive list"""
    job_lower = job_description.lower()
    
    skills = [
        # Programming Languages
        'javascript', 'python', 'java', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
        # Web Technologies
        'react', 'node.js', 'angular', 'vue', 'next.js', 'nuxt', 'express', 'django', 'flask', 'spring', 'laravel',
        'html', 'css', 'sass', 'scss', 'bootstrap', 'tailwind', 'jquery',
        # Databases
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'dynamodb', 'cassandra', 'elasticsearch',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible', 'jenkins', 'ci/cd', 'github actions',
        'git', 'gitlab', 'bitbucket', 'nginx', 'apache', 'linux', 'unix',
        # Data & AI
        'machine learning', 'data science', 'ai', 'artificial intelligence', 'tensorflow', 'pytorch', 'pandas', 'numpy',
        'tableau', 'power bi', 'spark', 'hadoop', 'kafka', 'rabbitmq',
        # Other Tools
        'graphql', 'rest api', 'microservices', 'agile', 'scrum', 'jira', 'confluence',
        # Soft Skills
        'leadership', 'communication', 'project management', 'teamwork', 'problem solving', 'analytics',
        'marketing', 'sales', 'design', 'writing', 'analysis'
    ]
    
    required = []
    for skill in skills:
        if skill in job_lower:
            required.append(skill.title())
    
    # Also extract any capitalized technical terms (likely tools/technologies)
    import re
    # Find capitalized words that might be tools/technologies
    tech_patterns = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', job_description)
    for pattern in tech_patterns:
        pattern_lower = pattern.lower()
        # If it's a known skill or looks like a technology name
        if pattern_lower in skills or (len(pattern) >= 3 and pattern not in required and pattern not in ['The', 'This', 'That', 'With', 'From']):
            if pattern not in required:
                required.append(pattern)
    
    return required[:20]  # Return top 20 required skills

def extract_resume_skills(resume_text: str) -> List[str]:
    """Extract skills mentioned in resume"""
    resume_lower = resume_text.lower()
    
    skills = [
        'javascript', 'python', 'java', 'react', 'node.js', 'sql', 'html', 'css',
        'typescript', 'angular', 'vue', 'express', 'django', 'flask', 'spring',
        'leadership', 'communication', 'project management', 'teamwork', 'problem solving',
        'marketing', 'sales', 'analytics', 'design', 'writing', 'analysis', 'agile',
        'devops', 'aws', 'docker', 'kubernetes', 'git', 'ci/cd', 'machine learning',
        'data science', 'ai', 'tensorflow', 'pytorch', 'postgresql', 'mongodb', 'redis',
        'elasticsearch', 'kafka', 'rabbitmq', 'nginx', 'apache', 'linux', 'terraform',
        'ansible', 'jenkins', 'github actions', 'graphql', 'rest api', 'microservices'
    ]
    
    found = []
    for skill in skills:
        if skill in resume_lower:
            found.append(skill.title())
    
    return found[:15]  # Return top 15 resume skills

def is_valid_skill(skill: str) -> bool:
    """Check if a string is a valid skill (not phone number, email, etc.)"""
    import re
    skill_lower = skill.lower().strip()
    
    # Filter out phone numbers (patterns like (501) 650-8445, 501-650-8445, etc.)
    phone_pattern = r'[\d\s\-\(\)]{10,}'
    if re.search(phone_pattern, skill):
        return False
    
    # Filter out emails
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    if re.search(email_pattern, skill):
        return False
    
    # Filter out URLs
    url_pattern = r'https?://|www\.'
    if re.search(url_pattern, skill_lower):
        return False
    
    # Filter out pure numbers or very short strings
    if len(skill_lower) < 2 or skill_lower.isdigit():
        return False
    
    # Filter out common non-skill words
    non_skills = ['com', 'www', 'http', 'https', 'email', 'phone', 'address', 'summary']
    if skill_lower in non_skills:
        return False
    
    return True

def extract_matched_skills(resume_text: str, job_description: str) -> List[str]:
    """Extract skills that match between resume and job description"""
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    
    skills = [
        'javascript', 'python', 'java', 'react', 'node.js', 'sql', 'html', 'css',
        'typescript', 'angular', 'vue', 'express', 'django', 'flask', 'spring',
        'leadership', 'communication', 'project management', 'teamwork', 'problem solving',
        'marketing', 'sales', 'analytics', 'design', 'writing', 'analysis', 'agile',
        'devops', 'aws', 'docker', 'kubernetes', 'git', 'ci/cd', 'machine learning',
        'data science', 'ai', 'tensorflow', 'pytorch', 'postgresql', 'mongodb', 'redis',
        'elasticsearch', 'kafka', 'rabbitmq', 'nginx', 'apache', 'linux', 'terraform',
        'ansible', 'jenkins', 'github actions', 'graphql', 'rest api', 'microservices'
    ]
    
    matched = []
    for skill in skills:
        if skill in job_lower and skill in resume_lower:
            skill_title = skill.title()
            if is_valid_skill(skill_title):
                matched.append(skill_title)
    
    return matched[:10]

def generate_improvement_suggestions(resume_text: str, job_description: str, match_score: int, skill_comparison: Dict, section_validation: Dict = None) -> List[str]:
    """Generate actionable suggestions to improve resume match score to 100%"""
    suggestions = []
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    
    # Get missing skills
    missing_skills = skill_comparison.get('missingSkills', [])
    job_required_skills = skill_comparison.get('jobRequiredSkills', [])
    resume_skills = skill_comparison.get('resumeSkills', [])
    
    # Add section validation suggestions if provided
    if section_validation:
        # Contact Info suggestions
        if section_validation["contactInfo"]["score"] < 100:
            if not section_validation["contactInfo"]["hasName"]:
                suggestions.append("Add your full name to the resume header")
            if not section_validation["contactInfo"]["hasPhone"]:
                suggestions.append("Include your phone number in the contact information section")
            if not section_validation["contactInfo"]["hasEmail"]:
                suggestions.append("Add your email address to the contact information section")
        
        # Summary suggestions
        if section_validation["summary"]["score"] < 100:
            if not section_validation["summary"]["hasSummary"]:
                suggestions.append("Add a professional summary section (2-3 lines describing who you are and what you can do)")
            elif not section_validation["summary"]["isProperLength"]:
                suggestions.append("Optimize your summary to be 2-3 lines (approximately 50-200 words)")
        
        # Work Experience suggestions
        if section_validation["workExperience"]["score"] < 100:
            if not section_validation["workExperience"]["hasExperience"]:
                suggestions.append("Add a work experience section with your past jobs")
            if not section_validation["workExperience"]["hasAchievements"]:
                suggestions.append("Include quantifiable achievements in your work experience (use numbers, percentages, metrics)")
            if not section_validation["workExperience"]["matchesJob"]:
                suggestions.append("Tailor your work experience descriptions to match keywords from the job description")
        
        # Skills suggestions
        if section_validation["skills"]["score"] < 100:
            if not section_validation["skills"]["hasSkills"]:
                suggestions.append("Add a skills section listing your technical and professional skills")
            if not section_validation["skills"]["matchesJobKeywords"]:
                suggestions.append("Update your skills section to include keywords from the job description")
        
        # Education suggestions
        if section_validation["education"]["score"] < 70:
            if not section_validation["education"]["hasEducation"]:
                suggestions.append("Include your education/degree information")
            if not section_validation["education"]["hasCertificates"]:
                suggestions.append("Add relevant certifications if you have them")
    
    # Calculate how many points each skill might add (rough estimate)
    points_needed = 100 - match_score
    skills_needed = max(1, min(len(missing_skills), points_needed // 10))
    
    # Suggest adding missing skills
    if missing_skills and skills_needed > 0:
        top_missing = missing_skills[:skills_needed]
        for skill in top_missing:
            suggestions.append(f"Add '{skill}' to your resume - this is a key requirement for this position")
    
    # Extract key requirements from job description
    key_phrases = []
    job_sentences = [s.strip() for s in job_lower.split('.') if len(s.strip()) > 30]
    
    # Look for requirement patterns
    requirement_keywords = ['required', 'must have', 'essential', 'necessary', 'should have', 'looking for']
    for sentence in job_sentences[:10]:
        if any(keyword in sentence for keyword in requirement_keywords):
            # Extract skills/technologies mentioned
            tech_keywords = ['python', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 'kubernetes', 
                           'experience', 'years', 'degree', 'certification', 'knowledge']
            if any(keyword in sentence for keyword in tech_keywords):
                key_phrases.append(sentence[:100])
    
    # Suggest adding experience/qualifications
    if 'years' in job_lower or 'experience' in job_lower:
        # Extract years requirement
        import re
        years_match = re.search(r'(\d+)\+?\s*years?', job_lower)
        if years_match:
            years_req = years_match.group(1)
            if f'{years_req} years' not in resume_lower and f'{years_req}+ years' not in resume_lower:
                suggestions.append(f"Highlight {years_req}+ years of relevant experience if you have it")
    
    # Suggest adding education if mentioned in job
    if ('degree' in job_lower or 'bachelor' in job_lower or 'master' in job_lower) and 'degree' not in resume_lower:
        suggestions.append("Include your educational qualifications if they match the job requirements")
    
    # Suggest adding certifications
    if 'certification' in job_lower or 'certified' in job_lower:
        if 'certification' not in resume_lower and 'certified' not in resume_lower:
            suggestions.append("Add relevant certifications if you have them - this is mentioned in the job requirements")
    
    # Suggest improving skill descriptions
    if len(resume_skills) < len(job_required_skills) * 0.7:
        suggestions.append("Expand your skills section to better match the job requirements")
    
    # Suggest adding relevant projects
    if 'project' in job_lower and 'project' not in resume_lower:
        suggestions.append("Add relevant projects or portfolio items that demonstrate your skills")
    
    # If score is very low, suggest major improvements
    if match_score < 50:
        suggestions.append("Consider gaining experience in the required technologies before applying")
        suggestions.append("Take online courses or certifications to bridge skill gaps")
    
    # Limit to top 5-7 most actionable suggestions
    return suggestions[:7] if suggestions else [
        "Review the job description carefully and add relevant keywords",
        "Highlight your most relevant experience at the top of your resume",
        "Ensure all required skills are mentioned in your resume"
    ]

def validate_resume_sections(resume_text: str, job_description: str, parsed_data: Dict, skill_comparison: Dict) -> Dict:
    """Validate all required resume sections according to professional standards"""
    import re
    
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    
    validation_results = {
        "contactInfo": {"hasName": False, "hasPhone": False, "hasEmail": False, "score": 0},
        "summary": {"hasSummary": False, "isProperLength": False, "score": 0},
        "workExperience": {"hasExperience": False, "hasAchievements": False, "matchesJob": False, "score": 0},
        "skills": {"hasSkills": False, "matchesJobKeywords": False, "score": 0},
        "education": {"hasEducation": False, "hasCertificates": False, "score": 0},
        "overallScore": 0
    }
    
    # 1. Contact Info Validation
    # Check for name
    name = parsed_data.get('name', '')
    if name and len(name.strip()) > 0:
        validation_results["contactInfo"]["hasName"] = True
    
    # Check for phone (patterns: (123) 456-7890, 123-456-7890, 123.456.7890, etc.)
    phone_patterns = [
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
        r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # International
        r'\d{10,}',  # 10+ digits
    ]
    for pattern in phone_patterns:
        if re.search(pattern, resume_text):
            validation_results["contactInfo"]["hasPhone"] = True
            break
    
    # Check for email
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    if re.search(email_pattern, resume_text):
        validation_results["contactInfo"]["hasEmail"] = True
    
    # Contact info score (all 3 required = 100%)
    contact_score = 0
    if validation_results["contactInfo"]["hasName"]:
        contact_score += 33.33
    if validation_results["contactInfo"]["hasPhone"]:
        contact_score += 33.33
    if validation_results["contactInfo"]["hasEmail"]:
        contact_score += 33.34
    validation_results["contactInfo"]["score"] = int(contact_score)
    
    # 2. Summary Validation (2-3 lines, about 50-200 words)
    summary_keywords = ['summary', 'profile', 'objective', 'about', 'overview']
    summary_section = None
    
    # Try to find summary section
    lines = resume_text.split('\n')
    in_summary = False
    summary_lines = []
    
    for i, line in enumerate(lines[:20]):  # Check first 20 lines
        line_lower = line.lower().strip()
        if any(keyword in line_lower for keyword in summary_keywords):
            in_summary = True
            continue
        if in_summary:
            if line.strip() and not any(keyword in line_lower for keyword in ['experience', 'education', 'skills', 'work']):
                summary_lines.append(line.strip())
            else:
                break
    
    if summary_lines:
        summary_text = ' '.join(summary_lines)
        word_count = len(summary_text.split())
        validation_results["summary"]["hasSummary"] = True
        # Check if it's 2-3 lines (approximately 50-200 words)
        if 50 <= word_count <= 200:
            validation_results["summary"]["isProperLength"] = True
            validation_results["summary"]["score"] = 100
        elif word_count > 0:
            validation_results["summary"]["score"] = 50  # Partial credit
    else:
        validation_results["summary"]["score"] = 0
    
    # 3. Work Experience Validation
    experience_keywords = ['experience', 'employment', 'work history', 'professional experience', 'career']
    has_experience = any(keyword in resume_lower for keyword in experience_keywords)
    validation_results["workExperience"]["hasExperience"] = has_experience
    
    # Check for achievements with numbers (metrics, percentages, etc.)
    achievement_patterns = [
        r'\d+%',  # Percentages
        r'\$\d+[KMB]?',  # Money amounts
        r'\d+\+',  # Numbers with +
        r'increased|decreased|improved|reduced|achieved|delivered|managed',
        r'\d+\s*(years?|months?)',  # Time periods
    ]
    has_achievements = False
    for pattern in achievement_patterns:
        if re.search(pattern, resume_lower, re.IGNORECASE):
            has_achievements = True
            break
    validation_results["workExperience"]["hasAchievements"] = has_achievements
    
    # Check if experience matches job description
    job_keywords = set(re.findall(r'\b\w{4,}\b', job_lower))  # Words 4+ chars
    resume_keywords = set(re.findall(r'\b\w{4,}\b', resume_lower))
    common_keywords = job_keywords.intersection(resume_keywords)
    match_ratio = len(common_keywords) / max(len(job_keywords), 1) if job_keywords else 0
    validation_results["workExperience"]["matchesJob"] = match_ratio >= 0.3  # 30% keyword overlap
    
    # Work experience score
    exp_score = 0
    if has_experience:
        exp_score += 40
    if has_achievements:
        exp_score += 30
    if validation_results["workExperience"]["matchesJob"]:
        exp_score += 30
    validation_results["workExperience"]["score"] = int(exp_score)
    
    # 4. Skills Validation
    matched_skills = skill_comparison.get("matchedSkills", [])
    job_required_skills = skill_comparison.get("jobRequiredSkills", [])
    
    validation_results["skills"]["hasSkills"] = len(skill_comparison.get("resumeSkills", [])) > 0
    validation_results["skills"]["matchesJobKeywords"] = len(matched_skills) > 0 and len(job_required_skills) > 0
    
    # Skills score based on match ratio
    if len(job_required_skills) > 0:
        skill_match_ratio = len(matched_skills) / len(job_required_skills)
        validation_results["skills"]["score"] = int(min(100, skill_match_ratio * 100))
    else:
        validation_results["skills"]["score"] = 50 if validation_results["skills"]["hasSkills"] else 0
    
    # 5. Education/Certificates Validation
    education_keywords = ['education', 'degree', 'bachelor', 'master', 'phd', 'university', 'college', 'diploma']
    cert_keywords = ['certification', 'certified', 'certificate', 'cert', 'license', 'licensed']
    
    has_education = any(keyword in resume_lower for keyword in education_keywords)
    has_certificates = any(keyword in resume_lower for keyword in cert_keywords)
    
    validation_results["education"]["hasEducation"] = has_education
    validation_results["education"]["hasCertificates"] = has_certificates
    
    # Education score
    edu_score = 0
    if has_education:
        edu_score += 70
    if has_certificates:
        edu_score += 30
    validation_results["education"]["score"] = int(edu_score)
    
    # Overall Score (weighted average)
    overall = (
        validation_results["contactInfo"]["score"] * 0.15 +      # 15% - Contact info
        validation_results["summary"]["score"] * 0.15 +          # 15% - Summary
        validation_results["workExperience"]["score"] * 0.30 +    # 30% - Work experience (most important)
        validation_results["skills"]["score"] * 0.25 +            # 25% - Skills
        validation_results["education"]["score"] * 0.15            # 15% - Education
    )
    validation_results["overallScore"] = int(overall)
    
    return validation_results

def check_ats_compatibility(resume_text: str, file_name: str) -> Dict:
    """Check ATS (Applicant Tracking System) compatibility"""
    import re
    
    ats_results = {
        "fileFormat": {"isATS": False, "format": "", "score": 0},
        "formatting": {"hasTables": False, "hasImages": False, "hasSpecialChars": False, "score": 0},
        "structure": {"hasHeaders": False, "isReadable": False, "score": 0},
        "overallScore": 0
    }
    
    # 1. File Format Check
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    if file_ext == 'pdf':
        ats_results["fileFormat"]["format"] = "PDF"
        ats_results["fileFormat"]["isATS"] = True
        ats_results["fileFormat"]["score"] = 100
    elif file_ext in ['docx', 'doc']:
        ats_results["fileFormat"]["format"] = "DOCX/DOC"
        ats_results["fileFormat"]["isATS"] = True
        ats_results["fileFormat"]["score"] = 90  # PDF is slightly better
    else:
        ats_results["fileFormat"]["format"] = file_ext.upper() if file_ext else "Unknown"
        ats_results["fileFormat"]["score"] = 0
    
    # 2. Formatting Issues
    # Check for tables (common ATS issue)
    table_indicators = ['|', '\t\t', '  +  ']  # Pipe, tabs, multiple spaces
    has_tables = any(indicator in resume_text for indicator in table_indicators)
    ats_results["formatting"]["hasTables"] = has_tables
    
    # Check for images (can't detect in text, but check for image references)
    image_keywords = ['[image]', '[img]', 'figure', 'photo']
    has_images = any(keyword in resume_text.lower() for keyword in image_keywords)
    ats_results["formatting"]["hasImages"] = has_images
    
    # Check for problematic special characters
    problematic_chars = ['•', '→', '←', '↑', '↓', '★', '☆', '◆', '■']
    has_special_chars = any(char in resume_text for char in problematic_chars)
    ats_results["formatting"]["hasSpecialChars"] = has_special_chars
    
    # Formatting score
    formatting_score = 100
    if has_tables:
        formatting_score -= 20
    if has_images:
        formatting_score -= 15
    if has_special_chars:
        formatting_score -= 10
    ats_results["formatting"]["score"] = max(0, formatting_score)
    
    # 3. Structure Check
    # Check for proper section headers
    section_headers = ['experience', 'education', 'skills', 'summary', 'objective', 'contact', 'work', 'employment']
    resume_lower = resume_text.lower()
    has_headers = any(header in resume_lower for header in section_headers)
    ats_results["structure"]["hasHeaders"] = has_headers
    
    # Check readability (proper line breaks, not all one block)
    lines = resume_text.split('\n')
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    avg_line_length = sum(len(line) for line in non_empty_lines) / max(len(non_empty_lines), 1)
    is_readable = 20 <= avg_line_length <= 100  # Reasonable line length
    ats_results["structure"]["isReadable"] = is_readable
    
    # Structure score
    structure_score = 0
    if has_headers:
        structure_score += 60
    if is_readable:
        structure_score += 40
    ats_results["structure"]["score"] = structure_score
    
    # Overall ATS Score
    overall = (
        ats_results["fileFormat"]["score"] * 0.30 +
        ats_results["formatting"]["score"] * 0.40 +
        ats_results["structure"]["score"] * 0.30
    )
    ats_results["overallScore"] = int(overall)
    
    return ats_results

def analyze_action_verbs(resume_text: str) -> Dict:
    """Analyze action verbs usage in resume"""
    import re
    
    # Strong action verbs (good)
    strong_verbs = [
        'achieved', 'accomplished', 'implemented', 'developed', 'created', 'designed',
        'managed', 'led', 'improved', 'increased', 'reduced', 'optimized', 'enhanced',
        'delivered', 'executed', 'established', 'launched', 'initiated', 'spearheaded',
        'transformed', 'streamlined', 'generated', 'produced', 'built', 'constructed',
        'analyzed', 'evaluated', 'researched', 'identified', 'resolved', 'solved',
        'collaborated', 'coordinated', 'facilitated', 'negotiated', 'presented', 'trained'
    ]
    
    # Weak verbs (should be replaced)
    weak_verbs = [
        'did', 'worked', 'helped', 'assisted', 'made', 'got', 'went', 'came',
        'tried', 'attempted', 'was', 'were', 'had', 'has', 'have', 'do', 'does'
    ]
    
    resume_lower = resume_text.lower()
    
    # Find all verbs (simple pattern matching)
    strong_found = []
    weak_found = []
    
    for verb in strong_verbs:
        pattern = r'\b' + re.escape(verb) + r'(ed|ing|s)?\b'
        matches = re.findall(pattern, resume_lower)
        if matches:
            strong_found.append(verb)
    
    for verb in weak_verbs:
        pattern = r'\b' + re.escape(verb) + r'(ed|ing|s)?\b'
        matches = re.findall(pattern, resume_lower)
        if matches:
            weak_found.append(verb)
    
    # Calculate score
    total_verbs = len(strong_found) + len(weak_found)
    if total_verbs == 0:
        verb_score = 50  # Neutral if no verbs detected
    else:
        strong_ratio = len(strong_found) / total_verbs
        verb_score = int(strong_ratio * 100)
    
    # Suggestions
    suggestions = []
    if len(weak_found) > 0:
        suggestions.append(f"Replace weak verbs like '{weak_found[0]}' with stronger action verbs")
    if len(strong_found) < 5:
        suggestions.append("Add more strong action verbs to make your achievements stand out")
    
    return {
        "strongVerbs": strong_found[:10],
        "weakVerbs": weak_found[:10],
        "strongCount": len(strong_found),
        "weakCount": len(weak_found),
        "score": verb_score,
        "suggestions": suggestions
    }

def analyze_keyword_density(resume_text: str, job_description: str) -> Dict:
    """Analyze keyword density and placement"""
    import re
    
    # Extract keywords from job description (important terms)
    job_lower = job_description.lower()
    job_words = re.findall(r'\b\w{4,}\b', job_lower)  # Words 4+ characters
    job_keyword_freq = {}
    for word in job_words:
        if word not in ['with', 'that', 'this', 'from', 'have', 'will', 'would', 'should']:
            job_keyword_freq[word] = job_keyword_freq.get(word, 0) + 1
    
    # Get top keywords from job description
    top_job_keywords = sorted(job_keyword_freq.items(), key=lambda x: x[1], reverse=True)[:20]
    top_keywords = [word for word, freq in top_job_keywords]
    
    # Check keyword presence in resume
    resume_lower = resume_text.lower()
    resume_words = set(re.findall(r'\b\w{4,}\b', resume_lower))
    
    found_keywords = []
    missing_keywords = []
    
    for keyword in top_keywords:
        if keyword in resume_words:
            found_keywords.append(keyword)
        else:
            missing_keywords.append(keyword)
    
    # Check keyword placement (top half of resume is better)
    resume_lines = resume_text.split('\n')
    top_half = ' '.join(resume_lines[:len(resume_lines)//2]).lower()
    keywords_in_top = sum(1 for kw in found_keywords if kw in top_half)
    
    # Calculate scores
    keyword_match_score = (len(found_keywords) / len(top_keywords) * 100) if top_keywords else 0
    placement_score = (keywords_in_top / len(found_keywords) * 100) if found_keywords else 0
    
    overall_score = int((keyword_match_score * 0.7) + (placement_score * 0.3))
    
    return {
        "foundKeywords": found_keywords[:15],
        "missingKeywords": missing_keywords[:15],
        "keywordsInTopHalf": keywords_in_top,
        "matchScore": int(keyword_match_score),
        "placementScore": int(placement_score),
        "overallScore": overall_score,
        "suggestions": [
            f"Add missing keywords: {', '.join(missing_keywords[:5])}" if missing_keywords else "Good keyword coverage",
            "Move important keywords to the top half of your resume" if placement_score < 50 else "Good keyword placement"
        ]
    }

def extract_quantifiable_achievements(resume_text: str) -> Dict:
    """Extract and analyze quantifiable achievements"""
    import re
    
    # Patterns for quantifiable achievements
    patterns = {
        "percentages": r'\d+%',
        "money": r'\$\d+[KMB]?|\d+\s*(million|billion|thousand|k|m|b)',
        "numbers": r'\d+[+\-]?\s*(users|customers|projects|employees|revenue|sales|deals|contracts)',
        "timeframes": r'\d+\s*(years?|months?|weeks?|days?)',
        "metrics": r'(increased|decreased|improved|reduced|achieved|delivered|managed|grew|saved)\s+\w+\s+by\s+\d+',
        "scales": r'\d+\s*(x|times|fold)'
    }
    
    achievements = {
        "percentages": [],
        "money": [],
        "numbers": [],
        "timeframes": [],
        "metrics": [],
        "scales": []
    }
    
    for category, pattern in patterns.items():
        matches = re.findall(pattern, resume_text, re.IGNORECASE)
        achievements[category] = matches[:5]  # Limit to 5 per category
    
    # Count total achievements
    total_achievements = sum(len(v) for v in achievements.values())
    
    # Calculate score
    if total_achievements >= 5:
        achievement_score = 100
    elif total_achievements >= 3:
        achievement_score = 75
    elif total_achievements >= 1:
        achievement_score = 50
    else:
        achievement_score = 0
    
    suggestions = []
    if total_achievements == 0:
        suggestions.append("Add quantifiable achievements with numbers, percentages, or metrics")
    elif total_achievements < 3:
        suggestions.append(f"Add more quantifiable achievements (currently {total_achievements}, aim for 5+)")
    
    return {
        "achievements": achievements,
        "totalCount": total_achievements,
        "score": achievement_score,
        "suggestions": suggestions
    }

def analyze_resume_length_structure(resume_text: str) -> Dict:
    """Analyze resume length and structure optimization"""
    import re
    
    # Calculate length metrics
    word_count = len(resume_text.split())
    char_count = len(resume_text)
    lines = resume_text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    
    # Estimate pages (assuming ~500 words per page)
    estimated_pages = word_count / 500
    
    # Length score
    if 400 <= word_count <= 1000:  # 1-2 pages ideal
        length_score = 100
    elif 200 <= word_count < 400:  # Too short
        length_score = 60
    elif 1000 < word_count <= 1500:  # Slightly long
        length_score = 80
    else:  # Very short or very long
        length_score = 40
    
    # Structure analysis
    # Check for bullet points
    bullet_patterns = [r'^[\-\*•]\s', r'^\d+\.\s', r'^[a-z]\)\s']
    has_bullets = any(re.search(pattern, resume_text, re.MULTILINE) for pattern in bullet_patterns)
    
    # Check section organization
    section_keywords = ['experience', 'education', 'skills', 'summary', 'objective', 'contact', 'work']
    sections_found = sum(1 for keyword in section_keywords if keyword in resume_text.lower())
    
    # Structure score
    structure_score = 0
    if has_bullets:
        structure_score += 50
    if sections_found >= 4:
        structure_score += 50
    elif sections_found >= 2:
        structure_score += 30
    
    suggestions = []
    if word_count < 400:
        suggestions.append(f"Resume is too short ({word_count} words). Aim for 400-1000 words (1-2 pages)")
    elif word_count > 1500:
        suggestions.append(f"Resume is too long ({word_count} words). Aim for 400-1000 words (1-2 pages)")
    
    if not has_bullets:
        suggestions.append("Use bullet points to make achievements easier to scan")
    
    if sections_found < 4:
        suggestions.append("Ensure you have clear sections: Experience, Education, Skills, Summary")
    
    return {
        "wordCount": word_count,
        "charCount": char_count,
        "estimatedPages": round(estimated_pages, 1),
        "hasBullets": has_bullets,
        "sectionsFound": sections_found,
        "lengthScore": length_score,
        "structureScore": structure_score,
        "overallScore": int((length_score * 0.6) + (structure_score * 0.4)),
        "suggestions": suggestions
    }

def calculate_modern_match_score(resume_text: str, job_description: str, skill_comparison: Dict, parsed_data: Dict = None) -> int:
    """Modern, comprehensive match score calculation with resume section validation"""
    if parsed_data is None:
        parsed_data = {}
    
    # Validate all resume sections
    section_validation = validate_resume_sections(resume_text, job_description, parsed_data, skill_comparison)
    
    # Get skill data
    matched_skills = skill_comparison.get("matchedSkills", [])
    missing_skills = skill_comparison.get("missingSkills", [])
    job_required_skills = skill_comparison.get("jobRequiredSkills", [])
    resume_skills = skill_comparison.get("resumeSkills", [])
    
    # Factor 1: Resume Section Completeness (30% weight) - Professional structure
    section_score = section_validation["overallScore"]
    
    # Factor 2: Skill Match Score (35% weight) - Most important for job match
    if len(job_required_skills) > 0:
        skill_match_ratio = len(matched_skills) / len(job_required_skills)
        skill_score = min(100, skill_match_ratio * 100)
    else:
        skill_score = 50  # Neutral if no specific skills mentioned
    
    # Factor 3: Semantic Similarity (25% weight) - Overall alignment
    semantic_score_raw = calculate_semantic_similarity(resume_text, job_description)
    semantic_score = min(100, max(0, semantic_score_raw * 100))
    
    # Factor 4: Completeness Score (10% weight) - Missing skills penalty
    if len(job_required_skills) > 0:
        missing_ratio = len(missing_skills) / len(job_required_skills)
        completeness_score = max(0, 100 - (missing_ratio * 100))
    else:
        completeness_score = 100
    
    # Weighted combination
    final_score = (
        section_score * 0.30 +        # 30% - Resume structure completeness
        skill_score * 0.35 +           # 35% - Skill matching (most important)
        semantic_score * 0.25 +         # 25% - Semantic similarity
        completeness_score * 0.10      # 10% - Completeness
    )
    
    # Normalize to 0-100 range
    match_score = int(max(0, min(100, final_score)))
    
    # CRITICAL: If all sections are perfect (100%) AND all skills match, give 100%
    all_sections_perfect = (
        section_validation["contactInfo"]["score"] == 100 and
        section_validation["summary"]["score"] == 100 and
        section_validation["workExperience"]["score"] == 100 and
        section_validation["skills"]["score"] == 100 and
        section_validation["education"]["score"] >= 70 and  # Education or certs
        len(missing_skills) == 0  # All required skills present
    )
    
    if all_sections_perfect:
        match_score = 100
        print("PERFECT RESUME: All sections complete and all skills match - Score: 100%")
    
    # Boost score if high skill match but low semantic (indicates good match)
    if skill_score >= 70 and semantic_score < 60:
        match_score = min(100, match_score + 5)
    
    # Reduce score if many missing critical skills
    if len(missing_skills) > len(matched_skills) and len(missing_skills) >= 3:
        match_score = max(0, match_score - 10)
    
    # Store section validation in skill_comparison for later use
    skill_comparison["sectionValidation"] = section_validation
    
    return match_score

def get_skill_comparison(resume_text: str, job_description: str) -> Dict:
    """Get detailed skill comparison between job and resume"""
    job_skills = extract_job_required_skills(job_description)
    resume_skills = extract_resume_skills(resume_text)
    matched_skills = extract_matched_skills(resume_text, job_description)
    
    # Calculate match percentage
    match_percentage = (len(matched_skills) / len(job_skills) * 100) if job_skills else 0
    
    # Get missing skills (required in job description but NOT in resume)
    # Compare job_skills with matched_skills to find what's missing
    missing_skills = []
    for job_skill in job_skills:
        # Check if this job skill is NOT in matched_skills
        if job_skill not in matched_skills:
            # Also check case-insensitive match
            job_skill_lower = job_skill.lower()
            is_matched = any(ms.lower() == job_skill_lower for ms in matched_skills)
            if not is_matched:
                missing_skills.append(job_skill)
    
    # Remove duplicates while preserving order
    seen = set()
    missing_skills_unique = []
    for skill in missing_skills:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            seen.add(skill_lower)
            missing_skills_unique.append(skill)
    missing_skills = missing_skills_unique
    
    # Get extra skills (in resume but not required)
    extra_skills = [s for s in resume_skills if s not in matched_skills]
    
    # Create skill scores with colors (green > yellow > orange > red)
    # Matched skills = highest (green), Extra skills = good (yellow), Missing skills = low (orange/red)
    skill_scores = []
    
    # Matched skills - Green (100%)
    for skill in matched_skills:
        skill_scores.append({
            "skill": skill,
            "score": 100,
            "status": "matched",
            "color": "#10B981"  # Green
        })
    
    # Extra skills - Yellow (70%)
    for skill in extra_skills[:10]:
        skill_scores.append({
            "skill": skill,
            "score": 70,
            "status": "extra",
            "color": "#F59E0B"  # Yellow/Amber
        })
    
    # Missing skills - Orange (40%) and Red (20%)
    for idx, skill in enumerate(missing_skills[:10]):
        # First half orange, second half red
        if idx < len(missing_skills) // 2:
            skill_scores.append({
                "skill": skill,
                "score": 40,
                "status": "missing",
                "color": "#F97316"  # Orange
            })
        else:
            skill_scores.append({
                "skill": skill,
                "score": 20,
                "status": "missing",
                "color": "#EF4444"  # Red
            })
    
    # Sort by score (highest first)
    skill_scores.sort(key=lambda x: x["score"], reverse=True)
    
    return {
        "jobRequiredSkills": job_skills,
        "resumeSkills": resume_skills,
        "matchedSkills": matched_skills,
        "missingSkills": missing_skills[:10],
        "extraSkills": extra_skills[:10],
        "matchPercentage": round(match_percentage, 1),
        "skillScores": skill_scores  # New: All skills with scores and colors
    }

def generate_strengths(resume_text: str, match_score: float, matched_skills: List[str], parsed_data: Dict) -> List[str]:
    """Generate strengths - Only based on matched skills from job description"""
    strengths = []
    
    # Only use matched skills that are required in job description and present in resume
    # Filter out any invalid skills (phone numbers, emails, etc.)
    valid_matched_skills = [s for s in matched_skills if is_valid_skill(s)]
    
    if valid_matched_skills:
        if len(valid_matched_skills) >= 1:
            strengths.append(f"Strong {valid_matched_skills[0]} skills")
        if len(valid_matched_skills) >= 2:
            strengths.append(f"Proficient in {valid_matched_skills[1]}")
        if len(valid_matched_skills) >= 3:
            strengths.append(f"Experienced with {valid_matched_skills[2]}")
    
    # Add general strengths only if we have matched skills
    if valid_matched_skills:
        if len(valid_matched_skills) >= 2:
            strengths.append("Multiple required skills present")
    
    return strengths[:3] if strengths else ["Some relevant skills identified"]

def generate_weaknesses(resume_text: str, job_description: str, match_score: float, matched_skills: List[str], missing_skills: List[str] = None) -> List[str]:
    """Generate weaknesses - ONLY show missing tools/skills from job description (NO generic messages)"""
    weaknesses = []
    
    # Get missing skills if not provided
    if missing_skills is None:
        skill_comparison = get_skill_comparison(resume_text, job_description)
        missing_skills = skill_comparison.get("missingSkills", [])
    
    # Ensure missing_skills is a list
    if not isinstance(missing_skills, list):
        missing_skills = []
    
    # Filter out invalid skills (phone numbers, emails, etc.)
    valid_missing_skills = [s for s in missing_skills if is_valid_skill(str(s))]
    
    # ONLY show specific missing tools/skills from job description - NO generic messages at all
    # Format: "Missing required skill/tool: [skill name]"
    for skill in valid_missing_skills[:15]:  # Show up to 15 missing skills
        skill_clean = str(skill).strip()
        if skill_clean:  # Only add non-empty skills
            weaknesses.append(f"Missing required skill/tool: {skill_clean}")
    
    # Return ONLY missing skills - NO generic messages like "Some skill gaps identified"
    return weaknesses

async def analyze_with_openai(resume_text: str, job_description: str, candidate_name: str, parsed_data: Dict = None) -> Dict:
    """Analyze resume using OpenAI GPT-4 API (Pure AI)"""
    if not OPENAI_AVAILABLE:
        raise ValueError("OpenAI library not installed. Install with: pip install openai")
    
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        raise ValueError("OPENAI_API_KEY not set. Cannot use OpenAI analysis.")
    
    client = openai.OpenAI(api_key=openai_api_key)
    
    prompt = f"""You are an expert resume analyzer. Analyze the following resume against the job description and provide a comprehensive assessment.

Job Description:
{job_description[:2000]}

Resume:
{resume_text[:3000]}

Provide a detailed analysis in JSON format with the following structure:
{{
  "matchScore": <number 0-100, representing overall match percentage>,
  "candidateName": "{candidate_name}",
  "strengths": ["strength1", "strength2", "strength3"],
  "weaknesses": ["weakness1", "weakness2", "weakness3"],
  "skillMatches": ["skill1", "skill2", "skill3", "skill4", "skill5"],
  "allSkills": ["skill1", "skill2", "skill3", ...],
  "skillComparison": {{
    "jobRequiredSkills": ["skill1", "skill2", ...],
    "resumeSkills": ["skill1", "skill2", ...],
    "matchedSkills": ["skill1", "skill2", ...],
    "missingSkills": ["skill1", "skill2", ...],
    "extraSkills": ["skill1", "skill2", ...],
    "matchPercentage": <number>,
    "skillScores": [
      {{"skill": "skill1", "score": 100, "status": "matched", "color": "#10B981"}},
      ...
    ]
  }}
}}

Respond ONLY with valid JSON, no markdown, no code blocks, no explanations."""

    try:
        response = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL", "gpt-4"),
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional resume analyzer. Always respond with valid JSON only, no markdown formatting, no code blocks."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000,
            response_format={"type": "json_object"}
        )
        
        response_text = response.choices[0].message.content
        # Clean response
        if response_text.startswith("```"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        
        analysis = json.loads(response_text)
        
        # Ensure all required fields exist
        ai_match_score = int(analysis.get("matchScore", 0))
        skill_comparison = analysis.get("skillComparison", {})
        
        # If skillComparison is missing or incomplete, generate it
        if not skill_comparison or not skill_comparison.get("missingSkills"):
            skill_comparison = get_skill_comparison(resume_text, job_description)
        
        # Recalculate using modern scoring method for accuracy
        if parsed_data is None:
            parsed_data = {"name": candidate_name}
        match_score = calculate_modern_match_score(resume_text, job_description, skill_comparison, parsed_data)
        
        # Use AI score as reference but prefer modern calculation
        # If AI score is very different (>20 points), use average for better accuracy
        if abs(ai_match_score - match_score) > 20:
            match_score = int((ai_match_score * 0.3) + (match_score * 0.7))
        
        # Generate weaknesses based on missing skills from job description
        missing_skills_list = skill_comparison.get("missingSkills", [])
        raw_skill_matches = analysis.get("skillMatches", [])
        filtered_skill_matches = [s for s in raw_skill_matches if is_valid_skill(str(s))][:5]
        
        # If no valid matched skills from AI, use skillComparison matched skills
        if not filtered_skill_matches and skill_comparison.get("matchedSkills"):
            filtered_skill_matches = [s for s in skill_comparison.get("matchedSkills", []) if is_valid_skill(str(s))][:5]
        
        # Generate weaknesses based on missing skills only
        weaknesses = generate_weaknesses(resume_text, job_description, match_score, filtered_skill_matches, missing_skills_list)
        
        # Generate improvement suggestions if score < 100%
        improvement_suggestions = analysis.get("improvementSuggestions", [])
        if not improvement_suggestions and match_score < 100:
            section_validation = skill_comparison.get("sectionValidation", {})
            improvement_suggestions = generate_improvement_suggestions(
                resume_text, job_description, match_score, skill_comparison, section_validation
            )
        
        # Run advanced analyses
        ats_analysis = check_ats_compatibility(resume_text, candidate_name + ".pdf")
        action_verb_analysis = analyze_action_verbs(resume_text)
        keyword_analysis = analyze_keyword_density(resume_text, job_description)
        achievements_analysis = extract_quantifiable_achievements(resume_text)
        length_structure_analysis = analyze_resume_length_structure(resume_text)
        
        result = {
            "candidateName": analysis.get("candidateName", candidate_name),
            "matchScore": match_score,
            "strengths": analysis.get("strengths", [])[:3],
            "weaknesses": weaknesses[:5],  # Use generated weaknesses based on missing skills
            "skillMatches": filtered_skill_matches,
            "allSkills": [s for s in analysis.get("allSkills", []) if is_valid_skill(str(s))][:10],
            "improvementSuggestions": improvement_suggestions,
            "skillComparison": skill_comparison,
            "advancedAnalysis": {
                "atsCompatibility": ats_analysis,
                "actionVerbs": action_verb_analysis,
                "keywordDensity": keyword_analysis,
                "quantifiableAchievements": achievements_analysis,
                "lengthStructure": length_structure_analysis
            }
        }
        
        print(f"AI Analysis Complete - Match Score: {result['matchScore']}%")
        return result
        
    except Exception as e:
        print(f"OpenAI API Error: {e}")
        raise

async def analyze_with_huggingface(resume_text: str, job_description: str, candidate_name: str) -> Dict:
    """Analyze resume using Hugging Face Inference API (Pure AI)"""
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        raise ValueError("HF_TOKEN not set. Cannot use Hugging Face Inference API.")
    
    # Use a better model for text generation - try Mistral or use a simpler approach
    # Using a text-to-text model that's more reliable
    model_id = os.getenv("HF_INFERENCE_MODEL", "mistralai/Mistral-7B-Instruct-v0.2")
    api_url = f"https://api-inference.huggingface.co/models/{model_id}"
    
    prompt = f"""<s>[INST] You are an expert resume analyzer. Analyze the resume against the job description and provide a JSON response.

Job Description:
{job_description[:1500]}

Resume:
{resume_text[:2000]}

Provide analysis in this exact JSON format:
{{
  "matchScore": 75,
  "strengths": ["Strong technical skills", "Relevant experience"],
  "weaknesses": ["Missing some required skills"],
  "skillMatches": ["Python", "JavaScript", "React"]
}}

Only return the JSON, no other text. [/INST]"""

    headers = {"Authorization": f"Bearer {hf_token}"}
    payload = {
        "inputs": prompt,
        "parameters": {
            "max_new_tokens": 500,
            "temperature": 0.3,
            "return_full_text": False
        }
    }
    
    try:
        print(f"Calling Hugging Face API: {model_id}")
        response = requests.post(api_url, headers=headers, json=payload, timeout=90)
        
        if response.status_code == 503:
            raise ValueError(f"Model {model_id} is loading. Please wait a moment and try again, or use a different model.")
        
        response.raise_for_status()
        result_data = response.json()
        
        # Handle different response formats
        if isinstance(result_data, list) and len(result_data) > 0:
            result_text = result_data[0].get("generated_text", "")
        elif isinstance(result_data, dict):
            result_text = result_data.get("generated_text", "")
        else:
            result_text = str(result_data)
        
        # Try to extract JSON from response
        json_match = re.search(r'\{[^{}]*"matchScore"[^{}]*\}', result_text, re.DOTALL)
        if not json_match:
            # Try broader match
            json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        
        if json_match:
            try:
                analysis = json.loads(json_match.group())
            except json.JSONDecodeError:
                # If JSON parsing fails, create a basic response
                print(f"Warning: Could not parse JSON from response: {result_text[:200]}")
                raise ValueError("AI response format is invalid. Please try again or configure OpenAI API.")
        else:
            raise ValueError(f"Could not extract JSON from AI response: {result_text[:200]}")
        
        # Ensure all required fields
        match_score = analysis.get("matchScore", 0)
        if isinstance(match_score, str):
            match_score = int(re.search(r'\d+', match_score).group()) if re.search(r'\d+', match_score) else 0
        
        # Get skill comparison for missing skills
        skill_comparison = analysis.get("skillComparison", {})
        if not skill_comparison or not skill_comparison.get("missingSkills"):
            # Generate skill comparison if not provided by AI
            skill_comparison = get_skill_comparison(resume_text, job_description)
        
        # Recalculate using modern scoring method for accuracy
        ai_match_score = int(match_score) if isinstance(match_score, (int, float)) else 0
        if isinstance(match_score, str):
            ai_match_score = int(re.search(r'\d+', match_score).group()) if re.search(r'\d+', match_score) else 0
        
        # Note: parsed_data not available in Hugging Face analysis, will use None
        match_score = calculate_modern_match_score(resume_text, job_description, skill_comparison, None)
        
        # Use AI score as reference but prefer modern calculation
        if abs(ai_match_score - match_score) > 20:
            match_score = int((ai_match_score * 0.3) + (match_score * 0.7))
        
        # Filter matched skills to remove phone numbers, emails, etc.
        raw_skill_matches = analysis.get("skillMatches", []) if isinstance(analysis.get("skillMatches"), list) else []
        filtered_skill_matches = [s for s in raw_skill_matches if is_valid_skill(str(s))][:5]
        
        # If no valid matched skills from AI, use skillComparison matched skills
        if not filtered_skill_matches and skill_comparison.get("matchedSkills"):
            filtered_skill_matches = [s for s in skill_comparison.get("matchedSkills", []) if is_valid_skill(str(s))][:5]
        
        # Generate weaknesses based on missing skills from job description
        missing_skills_list = skill_comparison.get("missingSkills", [])
        weaknesses = generate_weaknesses(resume_text, job_description, match_score, filtered_skill_matches, missing_skills_list)
        
        return {
            "candidateName": candidate_name,
            "matchScore": int(match_score),
            "strengths": analysis.get("strengths", [])[:3] if isinstance(analysis.get("strengths"), list) else [],
            "weaknesses": weaknesses[:5],  # Use generated weaknesses based on missing skills
            "skillMatches": filtered_skill_matches,
            "allSkills": [s for s in (analysis.get("skillMatches", []) if isinstance(analysis.get("skillMatches"), list) else []) if is_valid_skill(str(s))][:10],
            "skillComparison": skill_comparison
        }
    except requests.exceptions.RequestException as e:
        print(f"Hugging Face API Request Error: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response status: {e.response.status_code}")
            print(f"Response body: {e.response.text[:500]}")
        raise ValueError(f"Hugging Face API error: {str(e)}. Please check your HF_TOKEN or try configuring OpenAI API.")
    except Exception as e:
        print(f"Hugging Face API Error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise

async def analyze_with_semantic_similarity(resume_text: str, job_description: str, candidate_name: str) -> Dict:
    """Analyze using AI semantic similarity (SentenceTransformer) - Pure AI approach"""
    print("Using PURE AI semantic similarity analysis (SentenceTransformer)...")
    
    # Get skill comparison first
    skill_comparison = get_skill_comparison(resume_text, job_description)
    
    # Parse resume for section validation (basic parsing)
    parsed_data_basic = {"name": candidate_name}  # Minimal parsed data for semantic similarity
    
    # Calculate modern match score
    match_score = calculate_modern_match_score(resume_text, job_description, skill_comparison, parsed_data_basic)
    
    # Use AI embeddings to extract and compare skills (PURE AI)
    embedder = get_embedder()
    if embedder:
        # Extract key phrases/skills using AI embeddings
        # Split text into sentences and find most relevant ones
        resume_sentences = [s.strip() for s in resume_text.split('.') if len(s.strip()) > 20][:20]
        job_sentences = [s.strip() for s in job_description.split('.') if len(s.strip()) > 20][:20]
        
        if resume_sentences and job_sentences:
            # Get embeddings for all sentences
            resume_embeddings = embedder.encode(resume_sentences)
            job_embeddings = embedder.encode(job_sentences)
            
            # Find best matches using cosine similarity
            from sklearn.metrics.pairwise import cosine_similarity
            similarities = cosine_similarity(resume_embeddings, job_embeddings)
            
            # Extract top matching skills/phrases
            matched_indices = []
            for i, row in enumerate(similarities):
                max_sim_idx = row.argmax()
                if row[max_sim_idx] > 0.3:  # Threshold for relevance
                    matched_indices.append((i, max_sim_idx, row[max_sim_idx]))
            
            # Sort by similarity and extract top matches
            matched_indices.sort(key=lambda x: x[2], reverse=True)
            matched_skills = [resume_sentences[idx[0]][:50] for idx in matched_indices[:5]]
            all_skills = resume_sentences[:10]
        else:
            matched_skills = []
            all_skills = []
    else:
        # Fallback if embedder not available
        matched_skills = []
        all_skills = []
    
    # Generate strengths and weaknesses based on matched/missing skills
    matched_skills_list = skill_comparison.get("matchedSkills", [])
    missing_skills_list = skill_comparison.get("missingSkills", [])
    
    # Generate strengths based on matched skills
    strengths = []
    valid_matched_skills = [s for s in matched_skills_list if is_valid_skill(str(s))]
    if valid_matched_skills:
        if len(valid_matched_skills) >= 1:
            strengths.append(f"Strong {valid_matched_skills[0]} skills")
        if len(valid_matched_skills) >= 2:
            strengths.append(f"Proficient in {valid_matched_skills[1]}")
        if len(valid_matched_skills) >= 3:
            strengths.append(f"Experienced with {valid_matched_skills[2]}")
    else:
        strengths = ["Some relevant experience"] if match_score >= 60 else []
    
    # Generate weaknesses based on missing skills only
    weaknesses = generate_weaknesses(resume_text, job_description, match_score, valid_matched_skills, missing_skills_list)
    
    # Generate improvement suggestions if score < 100%
    improvement_suggestions = []
    if match_score < 100:
        improvement_suggestions = generate_improvement_suggestions(
            resume_text, job_description, match_score, skill_comparison
        )
    
    print(f"PURE AI Analysis Complete - Match Score: {match_score}% (from SentenceTransformer AI model)")
    
    # Generate improvement suggestions if score < 100%
    improvement_suggestions = []
    if match_score < 100:
        improvement_suggestions = generate_improvement_suggestions(
            resume_text, job_description, match_score, skill_comparison
        )
    
    # Filter matched skills to remove phone numbers, emails, etc.
    filtered_matched_skills = [s for s in matched_skills if is_valid_skill(str(s))][:5]
    filtered_all_skills = [s for s in all_skills if is_valid_skill(str(s))][:10]
    
    # If no valid matched skills, use skillComparison matched skills
    if not filtered_matched_skills and skill_comparison.get("matchedSkills"):
        filtered_matched_skills = [s for s in skill_comparison.get("matchedSkills", []) if is_valid_skill(str(s))][:5]
    
    # Run advanced analyses
    ats_analysis = check_ats_compatibility(resume_text, candidate_name + ".pdf")
    action_verb_analysis = analyze_action_verbs(resume_text)
    keyword_analysis = analyze_keyword_density(resume_text, job_description)
    achievements_analysis = extract_quantifiable_achievements(resume_text)
    length_structure_analysis = analyze_resume_length_structure(resume_text)
    
    return {
        "candidateName": candidate_name,
        "matchScore": match_score,
        "strengths": strengths[:3],
        "weaknesses": weaknesses[:3],
        "skillMatches": filtered_matched_skills,
        "allSkills": filtered_all_skills,
        "improvementSuggestions": improvement_suggestions,
        "skillComparison": skill_comparison,
        "advancedAnalysis": {
            "atsCompatibility": ats_analysis,
            "actionVerbs": action_verb_analysis,
            "keywordDensity": keyword_analysis,
            "quantifiableAchievements": achievements_analysis,
            "lengthStructure": length_structure_analysis
        }
    }

async def analyze_resume_file(
    file_content: bytes,
    file_name: str,
    job_description: str
) -> Dict:
    """Main analysis function - Uses AI only (OpenAI or Hugging Face)"""
    # Extract text
    resume_text = extract_resume_text(file_content, file_name)
    
    if not resume_text or len(resume_text.strip()) < 50:
        raise ValueError("Could not extract sufficient text from resume")
    
    # Parse resume with PyResparser for candidate name
    parsed_data = parse_resume_with_pyresparser(file_content, file_name)
    candidate_name = parsed_data.get('name') or file_name.replace('.pdf', '').replace('.docx', '').replace('.doc', '') or "Unknown Candidate"
    
    # Run advanced analyses
    print("Running advanced resume analyses...")
    ats_analysis = check_ats_compatibility(resume_text, file_name)
    action_verb_analysis = analyze_action_verbs(resume_text)
    keyword_analysis = analyze_keyword_density(resume_text, job_description)
    achievements_analysis = extract_quantifiable_achievements(resume_text)
    length_structure_analysis = analyze_resume_length_structure(resume_text)
    
    # Try AI analysis - OpenAI first, then Hugging Face, then semantic similarity (AI-based)
    use_ai_only = os.getenv("USE_AI_ONLY", "true").lower() == "true"
    
    if use_ai_only:
        print("Using AI-only analysis mode...")
        
        # Try OpenAI first
        openai_key = os.getenv("OPENAI_API_KEY")
        if openai_key and openai_key != "your-openai-api-key-here":
            try:
                print("Attempting OpenAI analysis...")
                return await analyze_with_openai(resume_text, job_description, candidate_name, parsed_data)
            except Exception as e:
                print(f"OpenAI analysis failed: {e}")
                print("Falling back to Hugging Face...")
        
        # Try Hugging Face Inference API
        hf_token = os.getenv("HF_TOKEN")
        if hf_token:
            try:
                print("Attempting Hugging Face Inference API analysis...")
                return await analyze_with_huggingface(resume_text, job_description, candidate_name)
            except Exception as e:
                print(f"Hugging Face Inference API failed: {e}")
                print("Falling back to AI semantic similarity...")
        
        # Fallback: Use AI semantic similarity (still AI-based, just different approach)
        try:
            print("Using AI semantic similarity analysis (SentenceTransformer)...")
            return await analyze_with_semantic_similarity(resume_text, job_description, candidate_name)
        except Exception as e:
            print(f"Semantic similarity analysis failed: {e}")
            raise ValueError("All AI analysis methods failed. Please check: 1) OPENAI_API_KEY, 2) HF_TOKEN, 3) SentenceTransformer model installation.")
    
    # Fallback to old hybrid method if AI is disabled
    print("Using hybrid analysis (AI + Python)...")
    resume_skills = extract_skills_with_spacy(resume_text)
    similarity_score = calculate_semantic_similarity(resume_text, job_description)
    matched_skills = extract_matched_skills(resume_text, job_description)
    if not matched_skills and resume_skills:
        matched_skills = resume_skills[:5]
    
    skill_comparison = get_skill_comparison(resume_text, job_description)
    
    # Use modern match score calculation with parsed_data
    match_score = calculate_modern_match_score(resume_text, job_description, skill_comparison, parsed_data)
    
    strengths = generate_strengths(resume_text, similarity_score, matched_skills, parsed_data)
    missing_skills = skill_comparison.get("missingSkills", [])
    weaknesses = generate_weaknesses(resume_text, job_description, similarity_score, matched_skills, missing_skills)
    
    # Generate improvement suggestions if score < 100%
    improvement_suggestions = []
    if match_score < 100:
        improvement_suggestions = generate_improvement_suggestions(
            resume_text, job_description, match_score, skill_comparison
        )
    
    # Filter matched skills to remove phone numbers, emails, etc.
    filtered_matched_skills = [s for s in matched_skills if is_valid_skill(str(s))][:5]
    filtered_resume_skills = [s for s in (resume_skills[:10] if resume_skills else []) if is_valid_skill(str(s))]
    
    # If no valid matched skills, use skillComparison matched skills
    if not filtered_matched_skills and skill_comparison.get("matchedSkills"):
        filtered_matched_skills = [s for s in skill_comparison.get("matchedSkills", []) if is_valid_skill(str(s))][:5]
    
    return {
        "candidateName": candidate_name,
        "matchScore": match_score,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "skillMatches": filtered_matched_skills,
        "allSkills": filtered_resume_skills,
        "improvementSuggestions": improvement_suggestions,
        "skillComparison": skill_comparison
    }

# ==================== RESUME REWRITING FUNCTIONS ====================

def parse_resume_sections(resume_text: str) -> Dict[str, str]:
    """Parse resume into structured sections"""
    sections = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "contact": ""
    }
    
    # Normalize text
    text = resume_text.strip()
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Common section headers
    section_patterns = {
        "summary": r"^(summary|objective|profile|about|overview|professional summary|executive summary)",
        "experience": r"^(experience|work experience|employment|professional experience|career history|work history)",
        "education": r"^(education|academic|qualifications|degrees|certifications|certificates)",
        "skills": r"^(skills|technical skills|core competencies|competencies|expertise|technologies)",
        "projects": r"^(projects|project experience|key projects|notable projects)"
    }
    
    current_section = None
    section_content = []
    
    for line in lines:
        line_lower = line.lower()
        matched = False
        
        # Check if line is a section header
        for section_name, pattern in section_patterns.items():
            if re.match(pattern, line_lower, re.IGNORECASE):
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content).strip()
                
                # Start new section
                current_section = section_name
                section_content = []
                matched = True
                break
        
        if not matched and current_section:
            section_content.append(line)
        elif not matched and not current_section:
            # Content before any section header (likely summary or contact)
            if len(line) < 100 and ('@' in line or re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', line)):
                sections["contact"] += line + '\n'
            elif not sections["summary"]:
                sections["summary"] += line + '\n'
    
    # Save last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content).strip()
    
    # Clean up sections
    for key in sections:
        sections[key] = sections[key].strip()
    
    return sections

async def rewrite_resume_with_ai(resume_text: str, job_description: str, parsed_sections: Dict[str, str]) -> str:
    """Rewrite resume using AI (OpenAI or Hugging Face)"""
    # Build structured prompt
    system_prompt = """You are an expert resume editor and ATS (Applicant Tracking System) optimization specialist. Your task is to rewrite resumes to be:
1. ATS-friendly (no tables, icons, or complex formatting)
2. Professional and clear
3. Optimized with strong action verbs
4. Enhanced with measurable achievements
5. Tailored to match the job description keywords
6. Well-structured with clear sections

Output a fully improved resume in clean text format with clear sections:
- Contact Information
- Summary
- Skills
- Experience (with bullet points using action verbs + metrics)
- Education
- Projects (if applicable)

Make sure to:
- Use strong action verbs (Led, Developed, Implemented, Achieved, etc.)
- Include quantifiable metrics where possible
- Match keywords from the job description
- Maintain professional tone
- Keep it concise and impactful"""

    user_prompt = f"""Rewrite the following resume to be ATS-friendly and optimized for this job description.

JOB DESCRIPTION:
{job_description}

ORIGINAL RESUME:
{resume_text}

RESUME SECTIONS (for reference):
{json.dumps(parsed_sections, indent=2)}

Please provide a fully rewritten, improved resume that:
1. Matches the job description keywords
2. Uses strong action verbs
3. Includes quantifiable achievements
4. Is ATS-friendly (plain text, no tables/icons)
5. Has clear section headers
6. Is professional and impactful

Output the complete rewritten resume:"""

    # Try OpenAI first
    openai_key = os.getenv("OPENAI_API_KEY")
    if openai_key and openai_key != "your-openai-api-key-here" and OPENAI_AVAILABLE:
        try:
            client = openai.OpenAI(api_key=openai_key)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=3000
            )
            rewritten = response.choices[0].message.content.strip()
            print("SUCCESS: Resume rewritten using OpenAI")
            return rewritten
        except Exception as e:
            print(f"OpenAI rewrite failed: {e}, trying Hugging Face...")
    
    # Try Hugging Face Inference API
    hf_token = os.getenv("HF_TOKEN")
    if hf_token:
        try:
            api_url = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2"
            headers = {"Authorization": f"Bearer {hf_token}"}
            
            prompt_text = f"{system_prompt}\n\n{user_prompt}"
            payload = {
                "inputs": prompt_text,
                "parameters": {
                    "max_new_tokens": 2000,
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }
            
            response = requests.post(api_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            result = response.json()
            
            if isinstance(result, list) and len(result) > 0:
                rewritten = result[0].get("generated_text", "").strip()
                print("SUCCESS: Resume rewritten using Hugging Face")
                return rewritten
            else:
                raise Exception("Unexpected response format from Hugging Face")
        except Exception as e:
            print(f"Hugging Face rewrite failed: {e}")
            raise Exception("AI resume rewriting failed. Please check OPENAI_API_KEY or HF_TOKEN configuration.")
    
    raise Exception("No AI service available. Please configure OPENAI_API_KEY or HF_TOKEN.")

def create_docx_from_text(text: str, filename: str = "improved_resume.docx") -> bytes:
    """Create a professionally formatted DOCX file from text"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Parse text into sections
        lines = text.split('\n')
        in_bullet_list = False
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if in_bullet_list:
                    in_bullet_list = False
                continue
            
            # Check if line is a section header
            is_header = (
                (line.isupper() and len(line) < 50 and len(line.split()) < 5) or
                any(header in line.lower() for header in [
                    'summary', 'experience', 'education', 'skills', 'projects', 
                    'contact', 'objective', 'profile', 'work experience', 
                    'professional experience', 'employment', 'certifications'
                ])
            )
            
            # Check if line is a bullet point
            is_bullet = (
                line.startswith('•') or line.startswith('-') or 
                line.startswith('*') or line.startswith('▪') or
                (len(line) > 2 and line[0].isdigit() and line[1] in ['.', ')'])
            )
            
            if is_header:
                # Add spacing before header (except first)
                if i > 0:
                    doc.add_paragraph()
                
                # Add section header with formatting
                p = doc.add_paragraph()
                run = p.add_run(line.upper())
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
                p.paragraph_format.space_after = Pt(6)
                in_bullet_list = False
            elif is_bullet:
                # Add bullet point
                p = doc.add_paragraph()
                # Remove bullet character and add formatted text
                clean_line = line.lstrip('•-*▪').strip()
                if clean_line and clean_line[0].isdigit() and len(clean_line) > 2:
                    clean_line = clean_line.split('.', 1)[-1].strip()
                    clean_line = clean_line.split(')', 1)[-1].strip()
                
                run = p.add_run(clean_line)
                run.font.size = Pt(11)
                p.style = 'List Bullet'
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(4)
                in_bullet_list = True
            else:
                # Regular content
                p = doc.add_paragraph(line)
                run = p.runs[0] if p.runs else p.add_run(line)
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(4)
                in_bullet_list = False
        
        # Save to bytes
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to create DOCX: {str(e)}")

def create_pdf_from_text(text: str) -> bytes:
    """Create a professionally formatted PDF from text"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib.colors import HexColor
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from xml.sax.saxutils import escape as xml_escape
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                               rightMargin=0.7*inch, leftMargin=0.7*inch,
                               topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=HexColor('#003366'),
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=4,
            leading=14
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=4,
            leading=14
        )
        
        # Parse and build PDF content
        story = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            
            # Check if line is a section header
            is_header = (
                (line.isupper() and len(line) < 50 and len(line.split()) < 5) or
                any(header in line.lower() for header in [
                    'summary', 'experience', 'education', 'skills', 'projects', 
                    'contact', 'objective', 'profile', 'work experience', 
                    'professional experience', 'employment', 'certifications'
                ])
            )
            
            # Check if line is a bullet point
            is_bullet = (
                line.startswith('•') or line.startswith('-') or 
                line.startswith('*') or line.startswith('▪') or
                (len(line) > 2 and line[0].isdigit() and line[1] in ['.', ')'])
            )
            
            if is_header:
                if i > 0:
                    story.append(Spacer(1, 8))
                # Escape XML special characters for reportlab
                escaped_line = xml_escape(line.upper())
                story.append(Paragraph(escaped_line, title_style))
                story.append(Spacer(1, 4))
            elif is_bullet:
                clean_line = line.lstrip('•-*▪').strip()
                if clean_line and clean_line[0].isdigit() and len(clean_line) > 2:
                    clean_line = clean_line.split('.', 1)[-1].strip()
                    clean_line = clean_line.split(')', 1)[-1].strip()
                # Escape XML special characters
                escaped_line = xml_escape(clean_line)
                story.append(Paragraph(f"• {escaped_line}", bullet_style))
            else:
                # Escape XML special characters
                escaped_line = xml_escape(line)
                story.append(Paragraph(escaped_line, normal_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        raise Exception("reportlab not installed. Install with: pip install reportlab")
    except Exception as e:
        raise Exception(f"Failed to create PDF: {str(e)}")

def create_markdown_from_text(text: str) -> str:
    """Convert text to Markdown format"""
    lines = text.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            markdown_lines.append('')
            continue
        
        # Check if line is a section header
        is_header = (
            (line.isupper() and len(line) < 50 and len(line.split()) < 5) or
            any(header in line.lower() for header in [
                'summary', 'experience', 'education', 'skills', 'projects', 
                'contact', 'objective', 'profile', 'work experience', 
                'professional experience', 'employment', 'certifications'
            ])
        )
        
        # Check if line is a bullet point
        is_bullet = (
            line.startswith('•') or line.startswith('-') or 
            line.startswith('*') or line.startswith('▪') or
            (len(line) > 2 and line[0].isdigit() and line[1] in ['.', ')'])
        )
        
        if is_header:
            markdown_lines.append(f"\n## {line}\n")
        elif is_bullet:
            clean_line = line.lstrip('•-*▪').strip()
            if clean_line and clean_line[0].isdigit() and len(clean_line) > 2:
                clean_line = clean_line.split('.', 1)[-1].strip()
                clean_line = clean_line.split(')', 1)[-1].strip()
            markdown_lines.append(f"- {clean_line}")
        else:
            markdown_lines.append(line)
    
    return '\n'.join(markdown_lines)

async def rewrite_resume(
    file_content: bytes,
    file_name: str,
    job_description: str
) -> Dict:
    """Main function to rewrite resume - Returns all formats and re-analysis"""
    # Extract resume text
    resume_text = extract_resume_text(file_content, file_name)
    
    if not resume_text or len(resume_text.strip()) < 50:
        raise ValueError("Could not extract sufficient text from resume")
    
    # Parse resume sections
    parsed_sections = parse_resume_sections(resume_text)
    
    # Rewrite with AI
    rewritten_text = await rewrite_resume_with_ai(resume_text, job_description, parsed_sections)
    
    # Generate all export formats
    docx_bytes = create_docx_from_text(rewritten_text)
    pdf_bytes = create_pdf_from_text(rewritten_text)
    markdown_text = create_markdown_from_text(rewritten_text)
    
    # Re-analyze the rewritten resume
    rewritten_analysis = None
    try:
        # Create a temporary file-like object for analysis
        rewritten_file_content = rewritten_text.encode('utf-8')
        rewritten_analysis = await analyze_resume_file(
            file_content=rewritten_file_content,
            file_name="improved_resume.txt",
            job_description=job_description
        )
    except Exception as e:
        print(f"Warning: Could not re-analyze rewritten resume: {e}")
    
    # Encode all formats to base64
    import base64
    base_name = file_name.replace('.pdf', '').replace('.docx', '').replace('.doc', '') + '_improved'
    
    return {
        "originalResume": resume_text,
        "rewrittenResume": rewritten_text,
        "originalScore": None,  # Will be set by caller if available
        "rewrittenScore": rewritten_analysis.get("matchScore") if rewritten_analysis else None,
        "scoreImprovement": None,  # Will be calculated in frontend
        "docxBase64": base64.b64encode(docx_bytes).decode('utf-8'),
        "pdfBase64": base64.b64encode(pdf_bytes).decode('utf-8'),
        "markdownText": markdown_text,
        "fileName": base_name,
        "rewrittenAnalysis": rewritten_analysis  # Full analysis for comparison
    }